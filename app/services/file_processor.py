import io
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract

from app.core.logging import get_logger
from app.utils.exceptions import ProcessingError

logger = get_logger(__name__)

#ToDo update to use Docling.

class FileProcessor:
    """Service for processing files and extracting text from them."""
    def __init__(self):
        self.supported_formats = {
            'pdf': self._extract_pdf_text,
            'docx': self._extract_docx_text,
            'png': self._extract_image_text,
            'jpg': self._extract_image_text,
            'jpeg': self._extract_image_text
        }

    async def extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from file content based on file type"""
        try:
            if file_type.lower() not in self.supported_formats:
                raise ProcessingError(f"Unsupported file type: {file_type}")

            extractor = self.supported_formats[file_type.lower()]
            text = await extractor(file_content)

            # Clean and normalize text
            cleaned_text = self._clean_text(text)

            logger.info(f"Extracted {len(cleaned_text)} characters from {file_type} file")
            return cleaned_text

        except Exception as e:
            logger.error(f"Failed to extract text from {file_type} file: {e}")
            raise ProcessingError(f"Text extraction failed: {str(e)}")

    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {e}")
                    continue

            return "\n\n".join(text_content)

        except Exception as e:
            raise ProcessingError(f"PDF text extraction failed: {str(e)}")

    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n\n".join(text_content)

        except Exception as e:
            raise ProcessingError(f"DOCX text extraction failed: {str(e)}")

    async def _extract_image_text(self, file_content: bytes) -> str:
                """Extract text from image using OCR"""
                try:
                    image = Image.open(io.BytesIO(file_content))

                    # Convert to RGB if necessary
                    if image.mode != 'RGB':
                        image = image.convert('RGB')

                    # Use Tesseract OCR to extract text
                    text = pytesseract.image_to_string(image, lang='eng')

                    return text

                except Exception as e:
                    raise ProcessingError(f"Image OCR failed: {str(e)}")

    def _clean_text(self, text: str) -> str:
                """Clean and normalize extracted text"""
                if not text:
                    return ""

                # Remove excessive whitespace
                text = re.sub(r'\s+', ' ', text)

                # Remove special characters but keep medical symbols
                text = re.sub(r'[^\w\s\-\+\=\(\)\[\]\{\}\.\,\;\:\!\?\'\"\%\°\±\≤\≥\→\←\↑\↓]', ' ', text)

                # Remove excessive newlines
                text = re.sub(r'\n\s*\n', '\n\n', text)

                # Strip leading/trailing whitespace
                text = text.strip()

                return text

    def chunk_text(
                    self,
                    text: str,
                    chunk_size: int = 1000,
                    overlap: int = 200,
                    preserve_sentences: bool = True
            ) -> List[Dict[str, Any]]:
                """Split text into chunks with semantic boundary detection"""
                if not text or len(text) < chunk_size:
                    return [{"text": text, "chunk_index": 0}] if text else []

                chunks = []

                if preserve_sentences:
                    chunks = self._chunk_by_sentences(text, chunk_size, overlap)
                else:
                    chunks = self._chunk_by_size(text, chunk_size, overlap)

                # Add metadata to chunks
                processed_chunks = []
                for i, chunk_text in enumerate(chunks):
                    processed_chunks.append({
                        "text": chunk_text.strip(),
                        "chunk_index": i,
                        "word_count": len(chunk_text.split()),
                        "char_count": len(chunk_text)
                    })

                logger.info(f"Created {len(processed_chunks)} chunks from text")
                return processed_chunks

    def _chunk_by_sentences(self, text: str, chunk_size: int, overlap: int) -> List[str]:
                """Chunk text by sentence boundaries"""
                # Split into sentences using multiple delimiters
                sentences = re.split(r'[.!?]+\s+', text)

                chunks = []
                current_chunk = ""

                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue

                    # Check if adding this sentence would exceed chunk size
                    potential_chunk = current_chunk + " " + sentence if current_chunk else sentence

                    if len(potential_chunk) <= chunk_size:
                        current_chunk = potential_chunk
                    else:
                        # Save current chunk if it has content
                        if current_chunk:
                            chunks.append(current_chunk)

                        # Start new chunk with overlap
                        if chunks and overlap > 0:
                            # Get last few words for overlap
                            words = current_chunk.split()
                            overlap_words = words[-min(overlap // 5, len(words)):]  # Approximate word overlap
                            current_chunk = " ".join(overlap_words) + " " + sentence
                        else:
                            current_chunk = sentence

                # Add final chunk
                if current_chunk:
                    chunks.append(current_chunk)

                return chunks

    def _chunk_by_size(self, text: str, chunk_size: int, overlap: int) -> List[str]:
                """Chunk text by character size with overlap"""
                chunks = []
                start = 0

                while start < len(text):
                    end = start + chunk_size

                    # If this isn't the last chunk, try to break at word boundary
                    if end < len(text):
                        # Look for the last space within the chunk
                        last_space = text.rfind(' ', start, end)
                        if last_space > start:
                            end = last_space

                    chunk = text[start:end].strip()
                    if chunk:
                        chunks.append(chunk)

                    # Move start position with overlap
                    start = max(start + 1, end - overlap)

                return chunks

    def detect_subject_tags(self, text: str) -> List[str]:
                """Detect medical subject tags from text content"""
                # Medical subject keywords mapping
                subject_keywords = {
                    'anatomy': ['anatomy', 'anatomical', 'structure', 'organ', 'tissue', 'muscle', 'bone', 'nerve'],
                    'physiology': ['physiology', 'function', 'mechanism', 'process', 'homeostasis', 'metabolism'],
                    'pathology': ['pathology', 'disease', 'disorder', 'syndrome', 'condition', 'abnormal', 'lesion'],
                    'pharmacology': ['drug', 'medication', 'pharmaceutical', 'dosage', 'therapy', 'treatment',
                                     'prescription'],
                    'cardiology': ['heart', 'cardiac', 'cardiovascular', 'blood pressure', 'artery', 'vein',
                                   'circulation'],
                    'neurology': ['brain', 'neurological', 'nervous system', 'neuron', 'cognitive', 'memory',
                                  'seizure'],
                    'respiratory': ['lung', 'respiratory', 'breathing', 'airway', 'oxygen', 'pneumonia', 'asthma'],
                    'endocrinology': ['hormone', 'endocrine', 'diabetes', 'thyroid', 'insulin', 'gland', 'metabolism'],
                    'immunology': ['immune', 'antibody', 'antigen', 'infection', 'vaccine', 'allergy', 'inflammation'],
                    'oncology': ['cancer', 'tumor', 'malignant', 'benign', 'chemotherapy', 'radiation', 'oncology']
                }

                detected_tags = set()
                text_lower = text.lower()

                for subject, keywords in subject_keywords.items():
                    for keyword in keywords:
                        if keyword in text_lower:
                            detected_tags.add(subject)
                            break  # Found one keyword for this subject, move to next

                return list(detected_tags)

        # Global file processor instance
file_processor = FileProcessor()